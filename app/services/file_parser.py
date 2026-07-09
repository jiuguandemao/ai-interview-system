import re
from pathlib import Path

from docx import Document as WordDocument
from pypdf import PdfReader


SUPPORTED_SUFFIXES = {".pdf", ".docx", ".txt", ".md"}


class FileParseError(ValueError):
    pass


def normalize_text(text: str) -> str:
    text = text.replace("\x00", "").replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    pages = [(page.extract_text() or "") for page in reader.pages]
    text = normalize_text("\n\n".join(pages))
    if len(text) < 30:
        raise FileParseError("PDF 几乎没有可提取文本，可能是扫描件，需要先做 OCR")
    return text


def extract_docx(path: Path) -> str:
    document = WordDocument(str(path))
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    return normalize_text("\n".join(parts))


def extract_text(path: str | Path) -> str:
    file_path = Path(path)
    suffix = file_path.suffix.lower()
    if suffix not in SUPPORTED_SUFFIXES:
        raise FileParseError(f"不支持的文件类型：{suffix}")
    if suffix == ".pdf":
        return extract_pdf(file_path)
    if suffix == ".docx":
        return extract_docx(file_path)
    return normalize_text(file_path.read_text(encoding="utf-8", errors="ignore"))
